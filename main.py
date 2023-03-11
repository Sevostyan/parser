from bs4 import BeautifulSoup
import requests
from docx import Document
import aspose.words as aw
from pdfminer.converter import TextConverter
from pdfminer.layout import LAParams
from pdfminer.pdfdocument import PDFDocument
from pdfminer.pdfinterp import PDFResourceManager, PDFPageInterpreter
from pdfminer.pdfpage import PDFPage
from pdfminer.pdfparser import PDFParser
import pdfminer
import requests
from io import StringIO

class PDFparser:
    def __init__(self):
        pass

    def parse(self, src):
        result = []
        output_string = StringIO()
        with open(src, 'rb') as in_file:
            parser = PDFParser(in_file)
            doc = PDFDocument(parser)
            rsrcmgr = PDFResourceManager()
            device = TextConverter(rsrcmgr, output_string, laparams=LAParams())
            interpreter = PDFPageInterpreter(rsrcmgr, device)
            for page in PDFPage.create_pages(doc):
                interpreter.process_page(page)
        result = output_string.getvalue()
        return result.split('\n\n')[:-1]


class HTMLParser:
    def __init__(self):
        pass

    def parse(self, text):
        return BeautifulSoup(text, "html.parser")

    def get_tree(self, src):
        result = []
        soup = BeautifulSoup(src, "html.parser")
        for tag in soup.find_all():
            tag_child = []
            for child in tag.children:
                if child.name != None:
                    tag_child.append(f'{child.name}')
            if len(tag_child) != 0:
                result.append({
                    'name': tag.name,
                    'children': tag_child,
                    'parent': tag.parent.name
                    })
        return result


class DOCXparser:
    def __init__(self):
        pass

    def parse(self, src):
        result = []
        document = Document(src)
        all_paras = document.paragraphs
        for para in all_paras:
            result.append(para.text)
        return result


class DOCParser:
    def __init__(self):
        pass

    def parse(self, src):
        result = []
        document = aw.Document(src)
        last_str = document.get_text().split('\r')[-3].split('\x0c')[0]
        result = document.get_text().split('\r')[1:-3]
        result.append(last_str)
        return result